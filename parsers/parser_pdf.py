# -*- encoding: utf-8 -*-

import os
import fitz
import time
import re
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from parsers.api import TranslateAPI
from parsers.db import DB
import tempfile

# 正则匹配参考文献
def is_reference(target):
    return re.match(r'references', target, re.I)


# 正则匹配图片标注
def is_figure(target):
    return re.match(r'fig\..\.', target, re.I)


class ParserPDF(object):
    def __init__(self, row_id, src_file, des_file, src_lang, des_lang):
        self.row_id = row_id
        self.src_file = src_file
        self.des_file = des_file
        self.src_lang = src_lang
        self.des_lang = des_lang

    def calculate_total_progress(self):
        total = 0
        cur_pdf = fitz.open(self.src_file)  # 待翻译的pdf
        i = 0
        for cur_page in cur_pdf:
            blks = cur_page.get_text_blocks(flags=4)  # read text blocks of input page
            flag = 0  # 记录当前的循
            reference_flag = 0  # 判断是否在参考文献之后
            blks.append((1, 2, 3, 6))
            content = ""
            for num in range(len(blks)):  # loop through the blocks
                if num == 0:
                    content = blks[0][4].replace("\n", " ")
                if num < len(blks) - 1:
                    # 两个块y轴距离很近的话，这里以1.0为界，这里判断当前数的右下角的坐标y值
                    if (abs(blks[num + 1][1] - blks[num][3]) <= 1.0 and abs(
                            blks[num + 1][1] - blks[num][3]) >= 0):
                        # 当前块在参考文献之后
                        if reference_flag == 1:
                            pass
                        # 其它情况
                        else:
                            flag = 1  #
                            content += blks[num + 1][4].replace("\n", " ")
                    else:
                        if flag == 1:
                            total += 1
                            flag = 0
                        else:
                            trans_paragraph = blks[num][4].replace("\n", " ")  # 将待翻译的句子换行换成空格
                            if is_reference(trans_paragraph.replace(' ', '')):
                                reference_flag = 1
                            else:
                                total += 1
                        try:
                            content = blks[num + 1][4].replace("\n", " ")
                        except:
                            pass
            i += 1

        return total

    def parse(self):
        total = self.calculate_total_progress()
        current = 0
        percent = 0
        trans = TranslateAPI()
        db = DB()
        root = tempfile.gettempdir()
        t0 = time.time()
        cur_pdf = fitz.open(self.src_file)  # 待翻译的pdf
        new_docx = Document()  # 翻译完成后要写入的docx
        new_docx.styles['Normal'].font.name = u'宋体'  # 设置翻译完成后的字体
        new_docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置翻译完成后的字体
        i = 0  # 定义页面数的递增
        for cur_page in cur_pdf:
            print('正在翻译第{}页...'.format(i + 1))
            img_list = cur_page.get_images()  # 获取当前页面的图片对象
            imgcount = 0
            for img in img_list:  # 获取当前页面的图像列表
                pix_temp1 = fitz.Pixmap(cur_pdf, img[0])
                if img[1]:
                    pix_temp2 = fitz.Pixmap(cur_pdf, img[1])
                    pix_temp = fitz.Pixmap(pix_temp1)
                    pix_temp.setAlpha(pix_temp2.samples)
                else:
                    pix_temp = pix_temp1
                imgcount += 1
                new_name = "图片{}.png".format(imgcount)  # 生成图片的名称
                pix_temp.save(os.path.join(root, new_name))
            blks = cur_page.get_text_blocks(flags=4)  # read text blocks of input page
            flag = 0  # 记录当前的循
            reference_flag = 0  # 判断是否在参考文献之后
            blks.append((1, 2, 3, 6))
            content = ""
            imgcount = 0
            for num in range(len(blks)):  # loop through the blocks
                # 如果是本页面最后一个块,直接结束,因为最后一个是方便计算自己添加的。
                if num == len(blks) - 1:
                    break
                # 如果这个块里放的是图像.
                if blks[num][-1] == 1:
                    imgcount += 1
                    try:
                        path_img = os.path.join(root, '图片{}.png'.format(imgcount))  # 当前页面第几个图片的位置
                        new_docx.add_picture(path_img, width=Inches(3))  # 设置图片保存的宽度
                        os.remove(path_img)  # 输入到新的pdf之后就移除
                    except:
                        pass
                    continue  # 跳过下面的插入翻译后文字的过程

                if num == 0:
                    content = blks[0][4].replace("\n", " ")
                # 矩形块，b[0]b[1]为左上角的坐标，b[2]b[3]为右下角的坐标
                r = fitz.Rect(blks[num][:4])
                # 如果不是倒数第一个块，则进入此循环
                if num < len(blks) - 1:
                    # 两个块y轴距离很近的话，这里以1.0为界，这里判断当前数的右下角的坐标y值
                    if (abs(blks[num + 1][1] - blks[num][3]) <= 1.0 and abs(
                            blks[num + 1][1] - blks[num][3]) >= 0):
                        # 当前块在参考文献之后
                        if reference_flag == 1:
                            pass
                        # 其它情况
                        else:
                            flag = 1  #
                            content += blks[num + 1][4].replace("\n", " ")
                    else:
                        if flag == 1:
                            res = trans.translate(self.src_lang, self.des_lang, content)
                            new_docx.add_paragraph(res)  # 添加到新的docx文档中
                            current = current + 1
                            if percent != int(current * 100 / total):
                                percent = int(current * 100 / total)
                                db.update_record_progress(self.row_id, percent)
                            flag = 0
                        else:
                            trans_paragraph = blks[num][4].replace("\n", " ")  # 将待翻译的句子换行换成空格
                            if is_reference(trans_paragraph.replace(' ', '')):
                                reference_flag = 1
                            else:
                                res = trans.translate(self.src_lang, self.des_lang, trans_paragraph)
                                new_docx.add_paragraph(res)
                                current = current + 1
                                if percent != int(current * 100 / total):
                                    percent = int(current * 100 / total)
                                    db.update_record_progress(self.row_id, percent)

                        try:
                            content = blks[num + 1][4].replace("\n", " ")
                        except:
                            pass
            i += 1

        # 文件保存
        new_docx.save(self.des_file)
        t1 = time.time()
        print("Total translation time: %g sec" % (t1 - t0))
        trans.close()
        db.close()
